from __future__ import annotations

from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "feishu_token_factory_docs"
ASSETS = OUT / "assets"

FONT_CJK = "/System/Library/Fonts/Hiragino Sans GB.ttc"
FONT_LATIN = "/System/Library/Fonts/Supplemental/Arial.ttf"

INK = "#162033"
MUTED = "#64748B"
LINE = "#CBD5E1"
BG = "#F6F8FB"
WHITE = "#FFFFFF"
BLUE = "#2563EB"
BLUE_LIGHT = "#EAF1FF"
TEAL = "#0F9F85"
TEAL_LIGHT = "#E7F8F3"
ORANGE = "#EA7A2F"
ORANGE_LIGHT = "#FFF1E7"
PURPLE = "#7657D5"
PURPLE_LIGHT = "#F1EDFF"
RED = "#D94B4B"
RED_LIGHT = "#FDECEC"
DARK = "#111827"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    index = 1 if bold else 0
    return ImageFont.truetype(FONT_CJK, size=size, index=index)


def latin_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(FONT_LATIN, size=size)


def text_width(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont) -> float:
    return draw.textbbox((0, 0), text, font=fnt)[2]


def wrap(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont, width: int) -> list[str]:
    lines: list[str] = []
    for raw in text.split("\n"):
        if not raw:
            lines.append("")
            continue
        current = ""
        for ch in raw:
            trial = current + ch
            if current and text_width(draw, trial, fnt) > width:
                lines.append(current)
                current = ch
            else:
                current = trial
        if current:
            lines.append(current)
    return lines


def draw_wrapped(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    fnt: ImageFont.FreeTypeFont,
    fill: str,
    width: int,
    spacing: int = 8,
    anchor: str | None = None,
) -> int:
    x, y = xy
    lines = wrap(draw, text, fnt, width)
    line_h = fnt.size + spacing
    for line in lines:
        draw.text((x, y), line, font=fnt, fill=fill, anchor=anchor)
        y += line_h
    return y


def canvas(title: str, subtitle: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (1600, 900), BG)
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((32, 28, 1568, 872), radius=24, fill=WHITE, outline="#E2E8F0", width=2)
    draw.text((72, 62), title, font=font(34, True), fill=INK)
    draw.text((72, 111), subtitle, font=font(19), fill=MUTED)
    draw.line((72, 154, 1528, 154), fill="#E5EAF2", width=2)
    return image, draw


def box(
    draw: ImageDraw.ImageDraw,
    rect: tuple[int, int, int, int],
    title: str,
    detail: str = "",
    fill: str = WHITE,
    outline: str = LINE,
    title_color: str = INK,
    radius: int = 14,
    badge: str | None = None,
) -> None:
    x1, y1, x2, y2 = rect
    draw.rounded_rectangle(rect, radius=radius, fill=fill, outline=outline, width=2)
    if badge:
        draw.rounded_rectangle((x1 + 16, y1 + 14, x1 + 58, y1 + 48), radius=9, fill=outline)
        draw.text((x1 + 37, y1 + 31), badge, font=font(17, True), fill=WHITE, anchor="mm")
        tx = x1 + 72
    else:
        tx = x1 + 20
    draw.text((tx, y1 + 18), title, font=font(21, True), fill=title_color)
    if detail:
        draw_wrapped(draw, (x1 + 20, y1 + 58), detail, font(16), MUTED, x2 - x1 - 40, spacing=7)


def pill(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, fill: str, color: str) -> int:
    x, y = xy
    fnt = font(15, True)
    w = int(text_width(draw, text, fnt)) + 30
    draw.rounded_rectangle((x, y, x + w, y + 34), radius=17, fill=fill)
    draw.text((x + w / 2, y + 17), text, font=fnt, fill=color, anchor="mm")
    return w


def arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    color: str = "#94A3B8",
    width: int = 4,
    label: str | None = None,
) -> None:
    x1, y1 = start
    x2, y2 = end
    draw.line((x1, y1, x2, y2), fill=color, width=width)
    if abs(x2 - x1) >= abs(y2 - y1):
        sign = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - sign * 14, y2 - 9), (x2 - sign * 14, y2 + 9)]
    else:
        sign = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - 9, y2 - sign * 14), (x2 + 9, y2 - sign * 14)]
    draw.polygon(points, fill=color)
    if label:
        mx, my = (x1 + x2) // 2, (y1 + y2) // 2
        bb = draw.textbbox((mx, my), label, font=font(14, True), anchor="mm")
        draw.rounded_rectangle((bb[0] - 8, bb[1] - 4, bb[2] + 8, bb[3] + 4), radius=8, fill=WHITE)
        draw.text((mx, my), label, font=font(14, True), fill=MUTED, anchor="mm")


def save(image: Image.Image, name: str) -> Path:
    path = ASSETS / name
    image.save(path, format="PNG", optimize=True)
    return path


def overall_architecture() -> Path:
    image, draw = canvas("Token 工厂：平台总架构", "面向 OpenRouter 类场景，业务控制面与高性能数据面分离")
    labels = [
        ("入口与体验", "开发者 / SDK  ·  用户控制台  ·  运营中台", BLUE, BLUE_LIGHT),
        ("边缘接入", "CDN / WAF  ·  L7 LB  ·  API Gateway", TEAL, TEAL_LIGHT),
        ("数据面", "Auth & Quota  →  Model Router  →  Provider Adapter  →  SSE Proxy", ORANGE, ORANGE_LIGHT),
        ("控制面", "Catalog  ·  Provider  ·  Deployment  ·  Route Policy  ·  Price / SKU  ·  Approval", PURPLE, PURPLE_LIGHT),
        ("事件与数据", "PostgreSQL（交易真相） · Redis（热状态） · Redpanda/Kafka（事件） · ClickHouse（分析）", RED, RED_LIGHT),
        ("供应与推理", "第三方模型 API  ·  Kubernetes  ·  vLLM / SGLang  ·  GPU 节点池", DARK, "#EEF1F5"),
    ]
    y = 188
    for i, (name, detail, color, light) in enumerate(labels):
        draw.rounded_rectangle((90, y, 1510, y + 88), radius=16, fill=light, outline=color, width=2)
        draw.rounded_rectangle((112, y + 20, 272, y + 68), radius=12, fill=color)
        draw.text((192, y + 44), name, font=font(20, True), fill=WHITE, anchor="mm")
        draw.text((310, y + 44), detail, font=font(18, True if i in (2, 3) else False), fill=INK, anchor="lm")
        if i < len(labels) - 1:
            arrow(draw, (800, y + 88), (800, y + 108), color="#9CAAC0", width=3)
        y += 108
    return save(image, "00_platform_architecture.png")


def request_flow() -> Path:
    image, draw = canvas("请求、路由与计费链路", "同步链路只做必要判定；用量、成本、账单通过幂等事件异步闭环")
    steps = [
        ("1", "接入", "OpenAI 兼容请求\n生成 request_id", BLUE, BLUE_LIGHT),
        ("2", "鉴权与额度", "Key / 租户 / 权益\nQPS · RPM · TPM", TEAL, TEAL_LIGHT),
        ("3", "智能路由", "健康 / 地域 / 合规\n价格 / 延迟 / 容量", ORANGE, ORANGE_LIGHT),
        ("4", "模型执行", "Provider API\n或 vLLM / SGLang", PURPLE, PURPLE_LIGHT),
        ("5", "流式响应", "SSE 转发\n统一错误与元数据", BLUE, BLUE_LIGHT),
    ]
    x = 72
    for i, (num, title, detail, color, light) in enumerate(steps):
        box(draw, (x, 210, x + 260, 390), title, detail, fill=light, outline=color, badge=num)
        if i < len(steps) - 1:
            arrow(draw, (x + 260, 300), (x + 292, 300), color="#7C8DA8", width=4)
        x += 292

    draw.rounded_rectangle((130, 500, 1470, 760), radius=18, fill="#111827", outline="#111827")
    draw.text((172, 535), "异步事件闭环", font=font(24, True), fill=WHITE)
    event_steps = [
        ("UsageRecorded", "token · cache\nlatency", BLUE),
        ("RatedCharge", "按冻结价格版本计费", TEAL),
        ("ProviderCost", "核算渠道与自建成本", ORANGE),
        ("LedgerEntry", "客户应收 / 渠道应付入账", PURPLE),
        ("Analytics", "经营、模型与 SLA 分析", RED),
    ]
    x = 170
    for i, (title, detail, color) in enumerate(event_steps):
        draw.rounded_rectangle((x, 600, x + 220, 704), radius=12, fill="#1F2937", outline=color, width=2)
        draw.text((x + 110, 626), title, font=font(17, True), fill=color, anchor="mm")
        draw_wrapped(draw, (x + 20, 656), detail, font(14), "#CBD5E1", 180, spacing=5)
        if i < len(event_steps) - 1:
            arrow(draw, (x + 220, 652), (x + 242, 652), color="#94A3B8", width=3)
        x += 248
    arrow(draw, (840, 390), (840, 492), color=RED, width=4, label="请求结束/中断均投递")
    return save(image, "01_request_billing_flow.png")


def business_loop() -> Path:
    image, draw = canvas("从资源到增长的运营闭环", "运营动作版本化、可审批、可灰度、可回滚，并由真实用量反哺决策")
    nodes = [
        (800, 230, "资源", "节点 · 模型 · 渠道", BLUE, BLUE_LIGHT),
        (1180, 390, "生产", "部署 · 评测 · 调度", TEAL, TEAL_LIGHT),
        (1040, 675, "经营", "SKU · 定价 · 上下架", ORANGE, ORANGE_LIGHT),
        (560, 675, "用户价值", "调用 · 体验 · 账单", PURPLE, PURPLE_LIGHT),
        (420, 390, "观测风控", "SLA · 成本 · 合规", RED, RED_LIGHT),
    ]
    for x, y, title, detail, color, light in nodes:
        draw.ellipse((x - 142, y - 82, x + 142, y + 82), fill=light, outline=color, width=4)
        draw.text((x, y - 18), title, font=font(25, True), fill=color, anchor="mm")
        draw.text((x, y + 25), detail, font=font(16), fill=INK, anchor="mm")
    centers = [(n[0], n[1]) for n in nodes]
    for i in range(len(centers)):
        x1, y1 = centers[i]
        x2, y2 = centers[(i + 1) % len(centers)]
        dx, dy = x2 - x1, y2 - y1
        length = max((dx * dx + dy * dy) ** 0.5, 1)
        sx, sy = x1 + int(dx / length * 150), y1 + int(dy / length * 90)
        ex, ey = x2 - int(dx / length * 150), y2 - int(dy / length * 90)
        arrow(draw, (sx, sy), (ex, ey), color="#8191AA", width=5)
    draw.ellipse((650, 360, 950, 610), fill="#111827", outline="#111827")
    draw.text((800, 432), "控制面", font=font(30, True), fill=WHITE, anchor="mm")
    draw.text((800, 483), "策略版本 + 审批 + 灰度", font=font(17), fill="#D5DCE8", anchor="mm")
    draw.text((800, 520), "数据回流 + 决策建议", font=font(17), fill="#D5DCE8", anchor="mm")
    return save(image, "02_business_loop.png")


def tech_stack() -> Path:
    image, draw = canvas("技术选型与核心优势", "优先选择团队可维护、可观测、可扩展的成熟技术；早期控制面保持模块化单体")
    rows = [
        ("数据面", "Go + net/http + Chi v5", "低延迟 · 标准 HTTP 生态 · SSE 转发稳定", BLUE, BLUE_LIGHT),
        ("控制面", "TypeScript + NestJS + Fastify", "业务迭代快 · 类型契约统一 · 模块边界清晰", PURPLE, PURPLE_LIGHT),
        ("前端", "React + Ant Design + Vite + TS", "统一组件体系 · 双 SPA 独立发布 · 易测试", ORANGE, ORANGE_LIGHT),
        ("交易数据", "PostgreSQL + Redis", "强一致交易真相 · 热状态与限流高效", TEAL, TEAL_LIGHT),
        ("事件分析", "Redpanda/Kafka + ClickHouse", "解耦主链路 · 高吞吐明细与经营分析", RED, RED_LIGHT),
        ("推理平台", "Kubernetes + vLLM/SGLang", "弹性 GPU 调度 · 开源模型吞吐最大化", DARK, "#EEF1F5"),
        ("可观测", "OpenTelemetry + Prometheus/Grafana", "跨网关、路由、渠道、账单的一条证据链", BLUE, BLUE_LIGHT),
    ]
    y = 185
    for i, (layer, choice, benefit, color, light) in enumerate(rows):
        draw.rounded_rectangle((78, y, 1522, y + 78), radius=14, fill=light if i % 2 == 0 else WHITE, outline="#DCE3ED", width=2)
        draw.rounded_rectangle((98, y + 16, 238, y + 62), radius=11, fill=color)
        draw.text((168, y + 39), layer, font=font(18, True), fill=WHITE, anchor="mm")
        draw.text((278, y + 39), choice, font=font(19, True), fill=INK, anchor="lm")
        draw.text((870, y + 39), benefit, font=font(17), fill=MUTED, anchor="lm")
        y += 90
    return save(image, "03_tech_stack_advantages.png")


def gateway_sequence() -> Path:
    image, draw = canvas("后端与 API Gateway：同步请求时序", "路由判定本地化、配置版本化、用量事件幂等化")
    actors = ["SDK / Client", "API Gateway", "Model Router", "Provider Adapter", "Event Bus"]
    xs = [140, 450, 770, 1090, 1410]
    colors = [BLUE, TEAL, ORANGE, PURPLE, RED]
    for x, label, color in zip(xs, actors, colors):
        draw.rounded_rectangle((x - 105, 190, x + 105, 242), radius=12, fill=color)
        draw.text((x, 216), label, font=font(17, True), fill=WHITE, anchor="mm")
        draw.line((x, 242, x, 820), fill="#D4DBE6", width=2)
    calls = [
        (0, 1, 300, "OpenAI-compatible request"),
        (1, 1, 365, "Key / quota / rate limit"),
        (1, 2, 430, "candidate set + policy version"),
        (2, 3, 495, "selected route + frozen quote"),
        (3, 0, 585, "SSE stream / normalized error"),
        (1, 4, 690, "UsageRecorded (idempotency_key)"),
        (4, 4, 755, "rating / cost / ledger"),
    ]
    for a, b, y, label in calls:
        if a == b:
            draw.line((xs[a], y, xs[a] + 90, y), fill=colors[a], width=4)
            draw.arc((xs[a] + 50, y, xs[a] + 130, y + 45), 270, 90, fill=colors[a], width=4)
            draw.line((xs[a] + 90, y + 45, xs[a], y + 45), fill=colors[a], width=4)
            draw.polygon([(xs[a], y + 45), (xs[a] + 14, y + 37), (xs[a] + 14, y + 53)], fill=colors[a])
            label_x = xs[a] - 180 if xs[a] > 1300 else xs[a] + 14
            draw.text((label_x, y - 26), label, font=font(14, True), fill=INK)
        else:
            arrow(draw, (xs[a], y), (xs[b], y), color=colors[a], width=4, label=label)
    return save(image, "10_gateway_sequence.png")


def frontend_architecture() -> Path:
    image, draw = canvas("前端架构：两个独立项目、两条交付链路", "用户端与运营中台独立仓库、独立构建、独立部署；公共能力通过版本化制品交付")
    box(draw, (70, 200, 520, 410), "用户端 Web 项目", "独立 Git 仓库 · 独立 CI/CD\nReact + Vite + Ant Design", fill=BLUE_LIGHT, outline=BLUE, badge="U")
    box(draw, (1080, 200, 1530, 410), "运营中台项目", "独立 Git 仓库 · 独立 CI/CD\nReact + Vite + Ant Design", fill=PURPLE_LIGHT, outline=PURPLE, badge="O")
    box(draw, (575, 205, 1025, 405), "版本化公共制品", "私有 npm 包 · OpenAPI Client\nAuth SDK · Design Tokens · Telemetry", fill=TEAL_LIGHT, outline=TEAL, badge="S")
    arrow(draw, (520, 305), (565, 305), color=BLUE)
    arrow(draw, (1080, 305), (1035, 305), color=PURPLE)
    box(draw, (150, 535, 700, 745), "Public / Console BFF", "面向用户的聚合 API\n隐藏内部模型、价格和渠道结构", fill=ORANGE_LIGHT, outline=ORANGE, badge="B")
    box(draw, (900, 535, 1450, 745), "Admin BFF / Control API", "细粒度 RBAC + 数据范围\n审批、灰度、审计与批量操作", fill=RED_LIGHT, outline=RED, badge="A")
    arrow(draw, (300, 410), (300, 525), color=BLUE, label="用户会话")
    arrow(draw, (1300, 410), (1300, 525), color=PURPLE, label="管理员会话")
    draw.text((800, 790), "数据：TanStack Query ｜ 状态：Zustand ｜ 表单：React Hook Form + Zod ｜ 图表：ECharts", font=font(17, True), fill=INK, anchor="mm")
    return save(image, "20_frontend_architecture.png")


def data_billing_flow() -> Path:
    image, draw = canvas("数据、计费与分析：一条可追溯证据链", "价格快照、原始用量、计价结果、总账分录分层保存，支持重放与对账")
    stages = [
        ("Quote Snapshot", "请求开始冻结\nprice_version / route", BLUE, BLUE_LIGHT),
        ("Usage Event", "tokens / cache / latency\nrequest_id", TEAL, TEAL_LIGHT),
        ("Rating", "客户 SKU 计价\n优惠 / 舍入 / 税", ORANGE, ORANGE_LIGHT),
        ("Costing", "渠道成本 / GPU 成本\n故障转移差价", PURPLE, PURPLE_LIGHT),
        ("Ledger", "应收 / 应付 / 退款\n不可变分录", RED, RED_LIGHT),
    ]
    x = 55
    for i, (title, detail, color, light) in enumerate(stages):
        box(draw, (x, 205, x + 270, 370), title, detail, fill=light, outline=color, badge=str(i + 1))
        if i < len(stages) - 1:
            arrow(draw, (x + 270, 287), (x + 292, 287), color="#7587A3", width=4)
        x += 305
    stores = [
        ("PostgreSQL", "租户、目录、价格、余额、发票、总账", TEAL),
        ("Redis", "权益、限流、预算、健康度和配置缓存", BLUE),
        ("Redpanda / Kafka", "Usage / Cost / Audit 等领域事件", ORANGE),
        ("ClickHouse", "调用明细、模型性能、成本和经营分析", PURPLE),
    ]
    x = 80
    for title, detail, color in stores:
        draw.rounded_rectangle((x, 520, x + 340, 725), radius=16, fill="#111827", outline=color, width=3)
        draw.text((x + 28, 552), title, font=font(21, True), fill=color)
        draw_wrapped(draw, (x + 28, 604), detail, font(16), "#D3DBE7", 284, spacing=8)
        x += 375
    arrow(draw, (800, 375), (800, 510), color=RED, width=4, label="事件 + 关联 ID")
    return save(image, "30_data_billing_flow.png")


def infra_topology() -> Path:
    image, draw = canvas("基础设施、推理与可观测：部署拓扑", "业务平面、推理平面和数据平面隔离；全链路由 OpenTelemetry 贯通")
    box(draw, (55, 200, 350, 355), "公网入口", "DNS · CDN · WAF\nL7 Load Balancer", fill=BLUE_LIGHT, outline=BLUE, badge="E")
    box(draw, (405, 190, 745, 370), "业务平面", "Gateway / Router\nControl API / BFF\nWorker / Scheduler", fill=TEAL_LIGHT, outline=TEAL, badge="A")
    box(draw, (800, 190, 1150, 370), "推理平面", "GPU Node Pool\nvLLM / SGLang\nAutoscaling / Warm Pool", fill=ORANGE_LIGHT, outline=ORANGE, badge="I")
    box(draw, (1205, 200, 1500, 355), "外部渠道", "OpenAI / Anthropic\n国内外模型服务商", fill=PURPLE_LIGHT, outline=PURPLE, badge="P")
    arrow(draw, (350, 278), (395, 278), color=BLUE)
    arrow(draw, (745, 278), (790, 278), color=TEAL)
    arrow(draw, (1150, 278), (1195, 278), color=ORANGE)
    box(draw, (130, 520, 535, 735), "数据平面", "PostgreSQL · Redis\nRedpanda/Kafka · ClickHouse\n对象存储 · KMS/Vault", fill="#EEF1F5", outline=DARK, badge="D")
    box(draw, (650, 520, 970, 735), "可观测", "OpenTelemetry\nPrometheus / Grafana\nLogs / Traces / Alerts", fill=RED_LIGHT, outline=RED, badge="O")
    box(draw, (1085, 520, 1490, 735), "交付治理", "Kubernetes · GitOps\n灰度发布 · 自动回滚\nNetworkPolicy · Secret 管理", fill=PURPLE_LIGHT, outline=PURPLE, badge="G")
    arrow(draw, (575, 370), (390, 510), color=TEAL, label="事务/事件")
    arrow(draw, (800, 370), (800, 510), color=RED, label="metrics / traces")
    arrow(draw, (1020, 370), (1285, 510), color=PURPLE, label="部署与策略")
    return save(image, "40_infra_topology.png")


def security_zones() -> Path:
    image, draw = canvas("安全、权限与治理：纵深防御", "身份、数据、密钥、变更和审计按租户与项目隔离")
    layers = [
        ("边界安全", "CDN / WAF · DDoS · TLS · Bot 防护", BLUE, BLUE_LIGHT),
        ("身份与授权", "OIDC / SSO · API Key 指纹 · RBAC + 数据范围 · MFA", TEAL, TEAL_LIGHT),
        ("运行时治理", "限流 · 预算 · 地域与合规过滤 · Provider 凭证隔离", ORANGE, ORANGE_LIGHT),
        ("数据保护", "KMS/Vault · 静态加密 · 脱敏 · 保留期 · 默认不保存 Prompt", PURPLE, PURPLE_LIGHT),
        ("变更与审计", "草稿 → 校验 → 影响模拟 → 审批 → 灰度 → 回滚；全量审计", RED, RED_LIGHT),
    ]
    y = 190
    margins = [70, 130, 190, 250, 310]
    for i, ((title, detail, color, light), margin) in enumerate(zip(layers, margins)):
        draw.rounded_rectangle((margin, y, 1600 - margin, y + 105), radius=18, fill=light, outline=color, width=3)
        draw.text((margin + 32, y + 30), f"{i + 1:02d}  {title}", font=font(21, True), fill=color)
        draw.text((margin + 280, y + 53), detail, font=font(17), fill=INK, anchor="lm")
        y += 120
    draw.rounded_rectangle((610, 790, 990, 840), radius=14, fill="#111827")
    draw.text((800, 815), "request_id + actor_id + version_id = 可追溯证据", font=font(17, True), fill=WHITE, anchor="mm")
    return save(image, "50_security_zones.png")


def roadmap() -> Path:
    image, draw = canvas("实施路线图：先打通交易闭环，再做智能经营", "每阶段都有可验收产物，避免在主链路未稳定前过早微服务化")
    phases = [
        ("P0 · 2 周", "工程底座", "双仓库 / 双流水线\n契约、鉴权、观测", BLUE, BLUE_LIGHT),
        ("P1 · 4-6 周", "MVP 闭环", "模型目录 / API Key\nGateway / 单渠道\n用量与预付费", TEAL, TEAL_LIGHT),
        ("P2 · 4-6 周", "多渠道与运营", "智能路由 / 价格版本\n上下架 / 审批\n账单与经营看板", ORANGE, ORANGE_LIGHT),
        ("P3 · 4-8 周", "自建推理", "GPU 节点池\nvLLM/SGLang\n评测、灰度、弹性", PURPLE, PURPLE_LIGHT),
        ("P4 · 持续", "规模化治理", "多地域 / 容灾\n成本优化 / 合规\n策略自动化", RED, RED_LIGHT),
    ]
    draw.line((120, 490, 1480, 490), fill="#9AA9BF", width=8)
    x_positions = [140, 470, 800, 1130, 1460]
    for i, ((phase, title, detail, color, light), x) in enumerate(zip(phases, x_positions)):
        draw.ellipse((x - 22, 468, x + 22, 512), fill=color, outline=WHITE, width=5)
        top = 210 if i % 2 == 0 else 545
        bottom = 430 if i % 2 == 0 else 765
        left = max(55, min(x - 145, 1290))
        box(draw, (left, top, left + 290, bottom), title, detail, fill=light, outline=color, badge=str(i))
        draw.text((left + 145, top - 30 if i % 2 == 0 else bottom + 28), phase, font=font(17, True), fill=color, anchor="mm")
        if i % 2 == 0:
            arrow(draw, (x, 430), (x, 462), color=color, width=3)
        else:
            arrow(draw, (x, 540), (x, 518), color=color, width=3)
    return save(image, "60_roadmap.png")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.lstrip("#"))
    tc_pr.append(shd)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(38, 47, 66)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.35
    for name, size, color in [("Title", 26, INK), ("Heading 1", 18, BLUE), ("Heading 2", 14, INK)]:
        st = styles[name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
        st.font.bold = True


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph(style="Title")
    p.add_run(title)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(subtitle)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(MUTED.lstrip("#"))


def add_picture(doc: Document, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(7.15))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in c.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MUTED.lstrip("#"))


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "EAF1FF")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(BLUE.lstrip("#"))
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def build_docs(paths: dict[str, Path]) -> None:
    docs: list[tuple[str, str, str, list[tuple[str, object]]]] = [
        (
            "Token工厂架构",
            "OpenRouter 类 AI 模型聚合平台技术总览",
            "这是一页式评审入口：先看平台分层和关键链路，再进入各专项子页。",
            [
                ("picture", (paths["overall"], "图 1：业务控制面与高性能数据面分离")),
                ("picture", (paths["request"], "图 2：同步请求与异步计费解耦")),
                ("picture", (paths["loop"], "图 3：资源、生产、经营、用户价值、观测风控形成闭环")),
                ("picture", (paths["tech"], "图 4：技术选型及其核心收益")),
                ("heading", "评审结论"),
                ("bullets", [
                    "数据面使用 Go + net/http + Chi v5，控制面使用 TypeScript + NestJS + Fastify；先保持模块化单体，按真实瓶颈拆分。",
                    "用户端与运营中台统一使用 React + Ant Design + Vite + TypeScript，但采用两个独立仓库和两条 CI/CD；公共能力通过版本化私有 npm 包与 OpenAPI 契约交付。",
                    "PostgreSQL 保存交易真相，Redis 保存热状态，Redpanda/Kafka 解耦事件，ClickHouse 承载明细分析。",
                    "OpenAI 兼容 API 是第一等产品能力；路由、价格、配额、审批与审计全部版本化。",
                    "自建模型通过 Kubernetes + vLLM/SGLang 接入，与第三方渠道统一抽象为 ProviderRoute。",
                ]),
            ],
        ),
        (
            "00 技术选型与优势",
            "为什么选择这套技术栈，以及何时需要演进",
            "选型原则：主链路低延迟、交易可追溯、业务迭代快、团队可维护。",
            [
                ("picture", (paths["tech"], "技术栈按数据面、控制面、前端、数据、推理与观测分层")),
                ("heading", "核心选择"),
                ("table", (["层级", "技术", "主要优势", "演进边界"], [
                    ["数据面", "Go + net/http + Chi v5", "标准 HTTP 生态、并发与 SSE 稳定", "流量或团队边界明确后再拆 Router / Adapter"],
                    ["控制面", "TypeScript + NestJS + Fastify", "类型契约统一、模块化、迭代快", "按账单、目录等独立负载拆服务"],
                    ["前端", "React + Ant Design + Vite + TS", "技术栈统一、两个项目独立演进", "用户端与中台独立仓库、独立部署"],
                    ["交易数据", "PostgreSQL + Redis", "强一致 + 高速热状态", "先分区与读副本，再评估分库"],
                    ["事件分析", "Redpanda/Kafka + ClickHouse", "削峰解耦、高吞吐分析", "消费组按领域独立扩容"],
                    ["推理", "Kubernetes + vLLM/SGLang", "GPU 弹性和高吞吐", "按模型族、地域拆节点池"],
                ])),
                ("heading", "不建议现在做的事"),
                ("bullets", [
                    "不在 MVP 阶段拆十几个微服务；跨服务事务和排障成本会先于收益出现。",
                    "不把 ClickHouse 当交易库，不把 Redis 当余额真相，不在网关同步写复杂账单。",
                    "不自研模型推理核心引擎，优先复用 vLLM/SGLang 并把精力放在平台能力。",
                ]),
            ],
        ),
        (
            "01 后端与 API Gateway",
            "高性能数据面、业务控制面和 Provider 适配器的职责边界",
            "目标是兼容 OpenAI API，同时做到可路由、可计量、可降级、可追溯。",
            [
                ("picture", (paths["gateway"], "Gateway 同步请求时序")),
                ("heading", "后端框架与工程组件"),
                ("table", (["层级", "选型", "落地职责"], [
                    ["数据面 HTTP", "Go + net/http + Chi v5", "OpenAI 兼容接口、SSE、鉴权、限流、路由与渠道代理"],
                    ["控制面 API", "TypeScript + NestJS + Fastify", "租户、目录、渠道、定价、审批、账单和后台权限"],
                    ["PostgreSQL", "pgx + sqlc", "显式 SQL、类型安全查询、交易与账本"],
                    ["Redis", "go-redis", "权益、预算、限流、配置和渠道健康缓存"],
                    ["事件总线", "franz-go", "UsageRecorded、成本、账单与审计事件"],
                    ["可观测", "OpenTelemetry + slog/zap", "Trace、Metrics、结构化日志和 request_id"],
                ])),
                ("heading", "数据面职责"),
                ("bullets", [
                    "Chi 作为薄路由层，保留标准 net/http Handler、Context 取消和 Middleware 生态；已有 Gin 工程经验时可等价替换。",
                    "解析 API Key 与租户上下文，执行权益、预算和 QPS/RPM/TPM 限流。",
                    "解析模型别名，按健康、地域、合规、价格、延迟和容量构建候选路由。",
                    "通过 Provider Adapter 统一请求/响应、SSE、错误码、token 用量和重试语义。",
                    "请求结束或中断都投递幂等 UsageRecorded；主链路不等待复杂计价。",
                ]),
                ("heading", "控制面模块"),
                ("table", (["模块", "职责", "关键实体"], [
                    ["Identity", "租户、组织、项目、Key", "tenant / project / api_key"],
                    ["Catalog", "模型目录、能力、别名", "service / model_capability"],
                    ["Provider", "渠道、凭证、部署", "provider / deployment / route"],
                    ["Routing", "策略、版本、灰度", "route_policy / policy_version"],
                    ["Commerce", "SKU、价格、权益、账单", "sku / price_version / ledger"],
                    ["Governance", "审批、审计、风控", "approval / audit_log / risk_rule"],
                ])),
            ],
        ),
        (
            "02 前端与体验架构",
            "用户端与运营中台双应用架构",
            "统一采用 React + Ant Design + Vite + TypeScript，但拆成两个独立项目、两个代码仓库和两条部署链路。",
            [
                ("picture", (paths["frontend"], "两个独立前端项目、版本化公共制品与 BFF 的关系")),
                ("heading", "项目与部署边界"),
                ("table", (["边界", "用户端 Web", "运营中台"], [
                    ["代码", "独立 Git 仓库", "独立 Git 仓库"],
                    ["交付", "独立 CI/CD、制品与回滚", "独立 CI/CD、制品与回滚"],
                    ["入口", "公网域名 + CDN/WAF", "管理域名 + SSO/VPN/IP 策略"],
                    ["配置", "独立环境变量与密钥", "独立环境变量与密钥"],
                    ["发布", "按用户产品节奏发布", "按运营需求和审批节奏发布"],
                ])),
                ("bullets", [
                    "两个项目不使用 workspace、源码软链接或相对路径依赖，任何一方都能单独克隆、构建、测试、发布和回滚。",
                    "OpenAPI Client、Auth SDK、Design Tokens 和 Telemetry 通过私有 npm Registry 发布并使用语义化版本锁定。",
                    "公共包升级由各项目自主发起，经过兼容性测试后再发布，禁止公共包变更触发两个项目强制同时上线。",
                ]),
                ("heading", "用户端"),
                ("bullets", [
                    "React + Vite + TypeScript：承载模型广场、模型详情与对比、API Key、用量、余额、账单和团队。",
                    "Ant Design 提供主题 Token、基础组件和无障碍交互；用户端通过定制主题保持产品化视觉，不直接照搬后台布局。",
                    "当前采用 SPA；公开模型目录通过 CDN 缓存与预渲染快照改善首屏和搜索收录，SEO 成为核心指标时再增加独立渲染层。",
                    "流式体验中心直接消费 SSE，展示 request_id、模型、token 和首 token 延迟。",
                ]),
                ("heading", "运营中台"),
                ("bullets", [
                    "React + Ant Design + Vite + TypeScript：资源、生产、经营、风控、财务和系统治理的高密度工作台。",
                    "使用 Ant Design Table、Form、Drawer、Modal、Tabs 和 ConfigProvider 建立统一交互，不引入第二套 UI 组件库。",
                    "页面权限、按钮权限与数据范围三层校验；前端校验只改善体验，后端仍强制授权。",
                    "复杂列表统一采用服务端分页、查询状态入 URL、批量操作回执和可恢复的错误状态。",
                ]),
                ("heading", "前端基础规范"),
                ("table", (["关注点", "推荐", "原因"], [
                    ["应用与构建", "React + Vite + TypeScript", "两个独立仓库、两条流水线和两套制品"],
                    ["UI 与主题", "Ant Design + Design Tokens", "组件完整、主题统一、适合工作台和管理后台"],
                    ["路由", "React Router", "路由守卫、懒加载和查询状态标准化"],
                    ["Server State", "TanStack Query", "缓存、重试、失效和并发请求一致"],
                    ["UI State", "Zustand", "轻量、局部、避免全局状态膨胀"],
                    ["表单", "React Hook Form + Zod", "运行时校验和 TS 类型对齐"],
                    ["契约", "OpenAPI 生成 Client", "减少接口漂移与手写 DTO"],
                    ["图表", "Apache ECharts", "经营、用量、成本和 SLA 图表能力完整"],
                    ["测试", "Vitest + Testing Library + Playwright", "覆盖组件、集成和关键链路"],
                ])),
            ],
        ),
        (
            "03 数据、计费与分析",
            "事件驱动的用量、成本、账单与经营分析",
            "账单的每个数字都能回到请求、路由、价格版本和原始用量。",
            [
                ("picture", (paths["data"], "从价格快照到总账与分析的证据链")),
                ("heading", "计费不变量"),
                ("bullets", [
                    "请求开始冻结 quote_snapshot，任何后续调价都不能改变这次请求的计价依据。",
                    "usage_event、rated_charge、provider_cost、ledger_entry 分层保存，支持重放但避免重复入账。",
                    "所有消费者使用 idempotency_key；乱序事件按 request_id 与版本号处理。",
                    "余额采用账本分录而非直接覆盖；退款、调账、赠送和冻结均是独立分录。",
                ]),
                ("heading", "数据存储分工"),
                ("table", (["存储", "保存内容", "禁止事项"], [
                    ["PostgreSQL", "交易、配置、余额、账单和总账", "不保存高频全量 Trace 明细"],
                    ["Redis", "限流、权益、预算与健康缓存", "不作为余额或价格真相"],
                    ["Redpanda/Kafka", "领域事件与异步任务", "不替代可查询的业务状态"],
                    ["ClickHouse", "调用明细、聚合指标与经营分析", "不承载强一致交易更新"],
                ])),
            ],
        ),
        (
            "04 基础设施、推理与可观测",
            "Kubernetes 部署、GPU 推理平台和全链路可观测",
            "业务平面、推理平面、数据平面隔离，配置和发布全程可回滚。",
            [
                ("picture", (paths["infra"], "Token 工厂部署拓扑")),
                ("heading", "推理生命周期"),
                ("bullets", [
                    "模型入库 → 评测 → 容量规划 → 部署 → 准入 → 灰度 → 在售监控 → 退市回收。",
                    "节点池按 GPU 型号、地域与租赁/自有属性分组；支持 warm pool 和优先级抢占策略。",
                    "vLLM/SGLang 暴露统一推理接口，平台负责编排、路由、计量和生命周期，不修改核心推理引擎。",
                ]),
                ("heading", "SLO 与观测"),
                ("table", (["层级", "关键指标", "典型告警"], [
                    ["入口", "RPS、4xx/5xx、连接数", "攻击、入口拥塞"],
                    ["网关", "鉴权延迟、限流、SSE 中断", "缓存失效、热 Key"],
                    ["路由", "候选数、切换率、回退率", "渠道集中故障"],
                    ["推理", "TTFT、TPS、队列、GPU 利用率", "OOM、队列过长"],
                    ["计费", "事件延迟、重复率、对账差异", "消费积压、账实不符"],
                ])),
            ],
        ),
        (
            "05 安全、权限与治理",
            "多租户、密钥、数据、变更和审计治理",
            "默认最小权限，敏感动作可审批，任何关键变更都能追溯。",
            [
                ("picture", (paths["security"], "纵深防御与证据链")),
                ("heading", "关键控制"),
                ("bullets", [
                    "用户 API Key 与 Provider 凭证分域管理；数据库仅保存密文、指纹和必要元数据。",
                    "RBAC 管理能力权限，数据范围约束租户/组织/项目，属性策略处理地域和合规条件。",
                    "默认不持久化完整 prompt/response；需要留存时启用脱敏、加密、保留期和删除策略。",
                    "调价、余额调整、Key 吊销、路由覆盖、节点关停均进入审批和审计。",
                ]),
                ("heading", "变更发布流程"),
                ("table", (["阶段", "系统动作", "失败处理"], [
                    ["草稿", "生成不可变版本", "可丢弃"],
                    ["校验", "语法、依赖、权限、合规检查", "阻止提交"],
                    ["影响模拟", "估算流量、成本和客户影响", "要求补充说明"],
                    ["审批", "按风险等级多级审批", "退回草稿"],
                    ["灰度", "按租户/比例/地域发布", "自动暂停或回滚"],
                    ["全量", "固化版本并写审计", "一键回到前版本"],
                ])),
            ],
        ),
        (
            "06 实施路线图",
            "从 MVP 到规模化平台的分阶段交付",
            "先完成真实可收费的调用闭环，再扩展多渠道、自建推理和智能经营。",
            [
                ("picture", (paths["roadmap"], "建议实施顺序与时间窗口")),
                ("heading", "阶段验收标准"),
                ("table", (["阶段", "验收结果", "退出条件"], [
                    ["P0 工程底座", "CI/CD、契约、鉴权、日志与 Trace 可用", "可稳定发布和回滚"],
                    ["P1 MVP", "单渠道完成 Key→调用→用量→扣费", "账单可追溯、SSE 稳定"],
                    ["P2 运营", "多渠道路由、定价、上下架、审批", "运营可独立完成日常动作"],
                    ["P3 自建推理", "GPU 节点、模型部署、评测、弹性", "吞吐与成本达到目标"],
                    ["P4 规模化", "多地域、容灾、合规和成本优化", "满足企业客户 SLA"],
                ])),
                ("heading", "首期团队建议"),
                ("bullets", [
                    "后端 3-4 人：2 人控制面/计费，1-2 人网关/路由/渠道适配。",
                    "前端 2 人：用户端与运营中台各 1 人；分别维护独立仓库，通过版本化制品共享设计规范和接口契约。",
                    "平台/算法工程 1-2 人：Kubernetes、GPU 推理、可观测和容量。",
                    "测试/产品/设计按项目节奏配置；计费与路由必须有专项测试负责人。",
                ]),
            ],
        ),
    ]

    for filename, title, subtitle, blocks in docs:
        doc = Document()
        style_doc(doc)
        add_title(doc, title, subtitle)
        p = doc.add_paragraph()
        run = p.add_run("定位：" + subtitle)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(TEAL.lstrip("#"))
        for kind, payload in blocks:
            if kind == "picture":
                path, caption = payload
                add_picture(doc, path, caption)
            elif kind == "heading":
                doc.add_heading(str(payload), level=1)
            elif kind == "bullets":
                add_bullets(doc, payload)
            elif kind == "table":
                headers, rows = payload
                add_table(doc, headers, rows)
        doc.save(OUT / f"{filename}.docx")


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    ASSETS.mkdir(parents=True, exist_ok=True)
    paths = {
        "overall": overall_architecture(),
        "request": request_flow(),
        "loop": business_loop(),
        "tech": tech_stack(),
        "gateway": gateway_sequence(),
        "frontend": frontend_architecture(),
        "data": data_billing_flow(),
        "infra": infra_topology(),
        "security": security_zones(),
        "roadmap": roadmap(),
    }
    build_docs(paths)
    print(f"Generated {len(list(OUT.glob('*.docx')))} documents in {OUT}")


if __name__ == "__main__":
    main()
